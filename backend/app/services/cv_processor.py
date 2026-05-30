import re
from pathlib import Path
from pypdf import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter

RAW_PATTERNS = {
    "summary": r"(summary|objective|about\s*me|profile|personal\s*statement)",
    "experience": r"(experience|work\s*history|employment|professional\s*experience|work\s*experience)",
    "education": r"(education|academic|qualification|degree|university|school)",
    "skills": r"(skills|technical\s*skills|core\s*competenc|technologies|tools|proficienc)",
    "projects": r"(projects|portfolio|personal\s*projects|academic\s*projects)",
    "certifications": r"(certifications?|licenses?|accreditations?|courses?|training)",
    "awards": r"(awards?|honors?|achievements?|recognition)",
    "publications": r"(publications?|papers?|research)",
    "languages": r"(languages?|linguistic)",
    "references": r"(references?|referees?)",
    "contact": r"(contact|email|phone|address|linkedin|github)",
}

SECTION_PATTERNS = {
    name: re.compile(rf"(?i)\b{pattern}\b") 
    for name, pattern in RAW_PATTERNS.items()
}


def extract_text_from_pdf(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    text_parts = [page.extract_text().strip() for page in reader.pages if page.extract_text()]
    return "\n\n".join(text_parts)


def extract_text_from_docx(file_path: Path) -> str:
    doc = DocxDocument(str(file_path))
    text_parts = [para.text.strip() for para in doc.paragraphs if para.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    return "\n".join(text_parts)


def extract_text(file_path: str) -> str:
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Please upload a PDF or DOCX file.")


def chunk_cv_text(full_text: str) -> tuple[list[dict], list[str]]:
    """
    Splits CV text into chunks and simultaneously aggregates detected sections.
    Returns: (chunks, list_of_detected_section_names)
    """
    lines = full_text.split("\n")
    sections = []
    current_section = {"name": "general", "lines": []}
    detected_section_names = set() 

    for line in lines:
        stripped = line.strip()
        if not stripped:
            current_section["lines"].append(line)
            continue

        detected = None
        if len(stripped) < 50: 
            for section_name, pattern in SECTION_PATTERNS.items():
                if pattern.search(stripped):
                    detected = section_name
                    detected_section_names.add(section_name)
                    break

        if detected and current_section["lines"]:
            sections.append(current_section)
            current_section = {"name": detected, "lines": [line]}
        else:
            current_section["lines"].append(line)

    if current_section["lines"]:
        sections.append(current_section)

    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=500,
        chunk_overlap=50,
        separators=["\n\n", "\n", ". ", " ", ""],
        length_function=len,
    )

    chunks = []
    chunk_index = 0

    for section in sections:
        section_text = "\n".join(section["lines"]).strip()
        if not section_text:
            continue

        for sub_chunk in text_splitter.split_text(section_text):
            if sub_chunk.strip():
                chunks.append({
                    "text": sub_chunk.strip(),
                    "metadata": {
                        "section": section["name"],
                        "chunk_index": chunk_index,
                        "source": "cv",
                    }
                })
                chunk_index += 1

    return chunks, list(detected_section_names)


def process_cv(file_path: str) -> dict:
    full_text = extract_text(file_path)

    if not full_text.strip():
        raise ValueError("Could not extract any text. Please check the file format.")

    chunks, sections_detected = chunk_cv_text(full_text)

    return {
        "full_text": full_text,
        "sections_detected": sections_detected,
        "chunks": chunks,
    }