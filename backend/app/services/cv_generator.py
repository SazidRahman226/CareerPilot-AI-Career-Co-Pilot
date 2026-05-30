"""
CareerPilot — CV Generator Service
======================================
Generates professional PDF and DOCX CVs from structured form data.
Uses ReportLab for PDF generation and python-docx for DOCX.
"""

import io
import logging
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm, inch
from reportlab.lib.colors import HexColor
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
    HRFlowable,
)
from docx import Document as DocxDocument
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from app.models.schemas import CVBuilderRequest

logger = logging.getLogger(__name__)

# ============================
#  Color Palette
# ============================
PRIMARY_COLOR = HexColor("#1a1a2e")
ACCENT_COLOR = HexColor("#4f46e5")
TEXT_COLOR = HexColor("#1f2937")
MUTED_COLOR = HexColor("#6b7280")


def generate_pdf(data: CVBuilderRequest) -> bytes:
    """
    Generate a professionally styled PDF CV from structured form data.

    Args:
        data: CVBuilderRequest with all CV sections.

    Returns:
        PDF file as bytes.
    """
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        topMargin=20 * mm,
        bottomMargin=20 * mm,
        leftMargin=20 * mm,
        rightMargin=20 * mm,
    )

    # --- Custom Styles ---
    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "CVName",
        parent=styles["Title"],
        fontSize=22,
        textColor=PRIMARY_COLOR,
        spaceAfter=4,
        alignment=TA_CENTER,
        fontName="Helvetica-Bold",
    )

    contact_style = ParagraphStyle(
        "CVContact",
        parent=styles["Normal"],
        fontSize=9,
        textColor=MUTED_COLOR,
        spaceAfter=12,
        alignment=TA_CENTER,
    )

    section_style = ParagraphStyle(
        "CVSection",
        parent=styles["Heading2"],
        fontSize=12,
        textColor=ACCENT_COLOR,
        spaceBefore=14,
        spaceAfter=6,
        fontName="Helvetica-Bold",
        borderPadding=(0, 0, 2, 0),
    )

    body_style = ParagraphStyle(
        "CVBody",
        parent=styles["Normal"],
        fontSize=10,
        textColor=TEXT_COLOR,
        spaceAfter=4,
        leading=14,
        alignment=TA_JUSTIFY,
    )

    sub_heading_style = ParagraphStyle(
        "CVSubHeading",
        parent=styles["Normal"],
        fontSize=11,
        textColor=PRIMARY_COLOR,
        spaceAfter=2,
        fontName="Helvetica-Bold",
    )

    date_style = ParagraphStyle(
        "CVDate",
        parent=styles["Normal"],
        fontSize=9,
        textColor=MUTED_COLOR,
        spaceAfter=2,
    )

    bullet_style = ParagraphStyle(
        "CVBullet",
        parent=styles["Normal"],
        fontSize=10,
        textColor=TEXT_COLOR,
        leftIndent=15,
        spaceAfter=2,
        leading=13,
        bulletIndent=5,
    )

    elements = []
    pi = data.personal_info

    # --- Header: Name ---
    if pi.full_name:
        elements.append(Paragraph(pi.full_name, name_style))

    # --- Contact Line ---
    contact_parts = []
    if pi.email:
        contact_parts.append(pi.email)
    if pi.phone:
        contact_parts.append(pi.phone)
    if pi.location:
        contact_parts.append(pi.location)
    if pi.linkedin:
        contact_parts.append(pi.linkedin)
    if pi.github:
        contact_parts.append(pi.github)
    if pi.website:
        contact_parts.append(pi.website)

    if contact_parts:
        elements.append(Paragraph(" • ".join(contact_parts), contact_style))

    # Separator line
    elements.append(HRFlowable(
        width="100%", thickness=1,
        color=HexColor("#e5e7eb"),
        spaceAfter=8,
    ))

    # --- Summary ---
    if data.summary:
        elements.append(Paragraph("PROFESSIONAL SUMMARY", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT_COLOR, spaceAfter=6))
        elements.append(Paragraph(data.summary, body_style))

    # --- Experience ---
    if data.experience:
        elements.append(Paragraph("WORK EXPERIENCE", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT_COLOR, spaceAfter=6))
        for exp in data.experience:
            title_line = f"<b>{exp.position}</b>"
            if exp.company:
                title_line += f" — {exp.company}"
            elements.append(Paragraph(title_line, sub_heading_style))

            date_parts = []
            if exp.start_date:
                date_parts.append(exp.start_date)
            if exp.current:
                date_parts.append("Present")
            elif exp.end_date:
                date_parts.append(exp.end_date)
            if exp.location:
                date_parts.append(exp.location)
            if date_parts:
                elements.append(Paragraph(" | ".join(date_parts), date_style))

            if exp.description:
                elements.append(Paragraph(exp.description, body_style))

            for highlight in exp.highlights:
                if highlight.strip():
                    elements.append(Paragraph(f"• {highlight}", bullet_style))
            elements.append(Spacer(1, 6))

    # --- Education ---
    if data.education:
        elements.append(Paragraph("EDUCATION", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT_COLOR, spaceAfter=6))
        for edu in data.education:
            title_line = f"<b>{edu.degree}</b>"
            if edu.field_of_study:
                title_line += f" in {edu.field_of_study}"
            elements.append(Paragraph(title_line, sub_heading_style))

            inst_parts = [edu.institution] if edu.institution else []
            if edu.start_date:
                inst_parts.append(f"{edu.start_date} – {edu.end_date or 'Present'}")
            if edu.gpa:
                inst_parts.append(f"GPA: {edu.gpa}")
            if inst_parts:
                elements.append(Paragraph(" | ".join(inst_parts), date_style))

            if edu.description:
                elements.append(Paragraph(edu.description, body_style))
            elements.append(Spacer(1, 4))

    # --- Skills ---
    if data.skills:
        elements.append(Paragraph("SKILLS", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT_COLOR, spaceAfter=6))
        elements.append(Paragraph(" • ".join(data.skills), body_style))

    # --- Projects ---
    if data.projects:
        elements.append(Paragraph("PROJECTS", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT_COLOR, spaceAfter=6))
        for proj in data.projects:
            title_line = f"<b>{proj.name}</b>"
            if proj.technologies:
                title_line += f" — <i>{proj.technologies}</i>"
            elements.append(Paragraph(title_line, sub_heading_style))
            if proj.description:
                elements.append(Paragraph(proj.description, body_style))
            if proj.url:
                elements.append(Paragraph(f"Link: {proj.url}", date_style))
            elements.append(Spacer(1, 4))

    # --- Certifications ---
    if data.certifications:
        elements.append(Paragraph("CERTIFICATIONS", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT_COLOR, spaceAfter=6))
        for cert in data.certifications:
            if cert.strip():
                elements.append(Paragraph(f"• {cert}", bullet_style))

    # --- Awards ---
    if data.awards:
        elements.append(Paragraph("AWARDS & ACHIEVEMENTS", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT_COLOR, spaceAfter=6))
        for award in data.awards:
            if award.strip():
                elements.append(Paragraph(f"• {award}", bullet_style))

    # --- Languages ---
    if data.languages:
        elements.append(Paragraph("LANGUAGES", section_style))
        elements.append(HRFlowable(width="100%", thickness=0.5, color=ACCENT_COLOR, spaceAfter=6))
        elements.append(Paragraph(" • ".join(data.languages), body_style))

    # Build PDF
    doc.build(elements)
    buffer.seek(0)
    return buffer.read()


def generate_docx(data: CVBuilderRequest) -> bytes:
    """
    Generate a professionally styled DOCX CV from structured form data.

    Args:
        data: CVBuilderRequest with all CV sections.

    Returns:
        DOCX file as bytes.
    """
    doc = DocxDocument()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    pi = data.personal_info
    accent_rgb = RGBColor(79, 70, 229)
    muted_rgb = RGBColor(107, 114, 128)
    dark_rgb = RGBColor(26, 26, 46)

    def add_section_header(text: str):
        """Add a styled section header."""
        p = doc.add_paragraph()
        p.space_before = Pt(14)
        p.space_after = Pt(4)
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = accent_rgb
        # Add a thin line under the header
        border_p = doc.add_paragraph()
        border_p.space_before = Pt(0)
        border_p.space_after = Pt(6)
        pPr = border_p._element.get_or_add_pPr()
        from docx.oxml.ns import qn
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "4",
            qn("w:space"): "1",
            qn("w:color"): "4F46E5",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    # --- Name ---
    if pi.full_name:
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_p.add_run(pi.full_name)
        name_run.bold = True
        name_run.font.size = Pt(22)
        name_run.font.color.rgb = dark_rgb

    # --- Contact ---
    contact_parts = []
    if pi.email:
        contact_parts.append(pi.email)
    if pi.phone:
        contact_parts.append(pi.phone)
    if pi.location:
        contact_parts.append(pi.location)
    if pi.linkedin:
        contact_parts.append(pi.linkedin)
    if pi.github:
        contact_parts.append(pi.github)
    if pi.website:
        contact_parts.append(pi.website)

    if contact_parts:
        contact_p = doc.add_paragraph()
        contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_p.space_after = Pt(12)
        contact_run = contact_p.add_run(" • ".join(contact_parts))
        contact_run.font.size = Pt(9)
        contact_run.font.color.rgb = muted_rgb

    # --- Summary ---
    if data.summary:
        add_section_header("PROFESSIONAL SUMMARY")
        p = doc.add_paragraph(data.summary)
        p.style.font.size = Pt(10)

    # --- Experience ---
    if data.experience:
        add_section_header("WORK EXPERIENCE")
        for exp in data.experience:
            # Position & Company
            p = doc.add_paragraph()
            run = p.add_run(exp.position)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = dark_rgb
            if exp.company:
                run2 = p.add_run(f" — {exp.company}")
                run2.font.size = Pt(11)

            # Date & Location
            date_parts = []
            if exp.start_date:
                date_parts.append(exp.start_date)
            if exp.current:
                date_parts.append("Present")
            elif exp.end_date:
                date_parts.append(exp.end_date)
            if exp.location:
                date_parts.append(exp.location)
            if date_parts:
                date_p = doc.add_paragraph(" | ".join(date_parts))
                date_p.space_before = Pt(0)
                date_p.space_after = Pt(2)
                for run in date_p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = muted_rgb

            if exp.description:
                desc_p = doc.add_paragraph(exp.description)
                desc_p.space_before = Pt(2)

            for highlight in exp.highlights:
                if highlight.strip():
                    doc.add_paragraph(highlight, style="List Bullet")

    # --- Education ---
    if data.education:
        add_section_header("EDUCATION")
        for edu in data.education:
            p = doc.add_paragraph()
            title = edu.degree
            if edu.field_of_study:
                title += f" in {edu.field_of_study}"
            run = p.add_run(title)
            run.bold = True
            run.font.size = Pt(11)
            run.font.color.rgb = dark_rgb

            inst_parts = [edu.institution] if edu.institution else []
            if edu.start_date:
                inst_parts.append(f"{edu.start_date} – {edu.end_date or 'Present'}")
            if edu.gpa:
                inst_parts.append(f"GPA: {edu.gpa}")
            if inst_parts:
                inst_p = doc.add_paragraph(" | ".join(inst_parts))
                for run in inst_p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = muted_rgb

            if edu.description:
                doc.add_paragraph(edu.description)

    # --- Skills ---
    if data.skills:
        add_section_header("SKILLS")
        doc.add_paragraph(" • ".join(data.skills))

    # --- Projects ---
    if data.projects:
        add_section_header("PROJECTS")
        for proj in data.projects:
            p = doc.add_paragraph()
            run = p.add_run(proj.name)
            run.bold = True
            run.font.size = Pt(11)
            if proj.technologies:
                tech_run = p.add_run(f" — {proj.technologies}")
                tech_run.italic = True
                tech_run.font.size = Pt(10)
            if proj.description:
                doc.add_paragraph(proj.description)
            if proj.url:
                url_p = doc.add_paragraph(f"Link: {proj.url}")
                for run in url_p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = muted_rgb

    # --- Certifications ---
    if data.certifications:
        add_section_header("CERTIFICATIONS")
        for cert in data.certifications:
            if cert.strip():
                doc.add_paragraph(cert, style="List Bullet")

    # --- Awards ---
    if data.awards:
        add_section_header("AWARDS & ACHIEVEMENTS")
        for award in data.awards:
            if award.strip():
                doc.add_paragraph(award, style="List Bullet")

    # --- Languages ---
    if data.languages:
        add_section_header("LANGUAGES")
        doc.add_paragraph(" • ".join(data.languages))

    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
